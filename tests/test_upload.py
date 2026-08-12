import io
from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document

from app.main import create_app


def test_upload_docx_extracts_text(tmp_path: Path) -> None:
    app = create_app(upload_dir=tmp_path)
    client = TestClient(app)

    document = Document()
    document.add_paragraph("Hello from the AI knowledge assistant")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/upload",
        files={
            "file": (
                "sample.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "sample.docx"
    assert "Hello from the AI knowledge assistant" in payload["extracted_text"]


def test_rejects_unsupported_file_type(tmp_path: Path) -> None:
    app = create_app(upload_dir=tmp_path)
    client = TestClient(app)

    response = client.post(
        "/upload",
        files={"file": ("notes.txt", b"not supported", "text/plain")},
    )

    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


def test_lists_uploaded_documents(tmp_path: Path) -> None:
    app = create_app(upload_dir=tmp_path)
    client = TestClient(app)

    file_path = tmp_path / "example.docx"
    file_path.write_bytes(b"fake docx content")

    response = client.get("/documents")

    assert response.status_code == 200
    payload = response.json()
    assert any(item["filename"] == "example.docx" for item in payload)


def test_deletes_uploaded_document(tmp_path: Path) -> None:
    app = create_app(upload_dir=tmp_path)
    client = TestClient(app)

    file_path = tmp_path / "example.docx"
    file_path.write_bytes(b"fake docx content")

    response = client.delete("/documents/example.docx")

    assert response.status_code == 200
    assert response.json()["deleted"] is True
    assert not file_path.exists()


def test_upload_classifies_document_category(tmp_path: Path) -> None:
    app = create_app(upload_dir=tmp_path)
    client = TestClient(app)

    document = Document()
    document.add_paragraph("Employee benefits policy and annual leave for staff training")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/upload",
        files={
            "file": (
                "hr-policy.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["category"] == "HR"
    assert payload["confidence"] >= 0.5
